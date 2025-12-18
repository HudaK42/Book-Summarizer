import os
import PyPDF2
import pdfplumber
import docx

def extract_text_from_txt(file_path):
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
    except UnicodeDecodeError:
        try:
            with open(file_path, "r", encoding="latin-1") as f:
                text = f.read()
        except Exception as e:
            raise ValueError(f"Failed to read TXT file: {str(e)}")
    except Exception as e:
        raise ValueError(f"Unexpected error reading TXT file: {str(e)}")
    return text

def extract_text_from_pdf(file_path):
    text = ""
    try:
        pdf = PyPDF2.PdfReader(file_path)
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"

        if not text.strip():  # fallback to pdfplumber
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"

        if not text.strip():
            raise ValueError("PDF contains no extractable text (maybe scanned). OCR required.")
    except Exception as e:
        raise ValueError(f"Failed to extract PDF text: {str(e)}")
    return text

def extract_text_from_docx(file_path):
    try:
        doc = docx.Document(file_path)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\t"
                text += "\n"
    except Exception as e:
        raise ValueError(f"Failed to extract DOCX text: {str(e)}")
    return text

def extract_text(file_path):
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File does not exist: {file_path}")

    ext = file_path.rsplit(".", 1)[1].lower()
    try:
        if ext == "txt":
            text = extract_text_from_txt(file_path)
        elif ext == "pdf":
            text = extract_text_from_pdf(file_path)
        elif ext == "docx":
            text = extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")
    except Exception as e:
        raise ValueError(f"Error during text extraction: {str(e)}")

    # Clean & normalize
    try:
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        lines = [line.strip() for line in text.split("\n") if line.strip()]
        cleaned_text = "\n".join(lines)
    except Exception as e:
        raise ValueError(f"Error during text cleaning: {str(e)}")

    if not cleaned_text:
        raise ValueError("Extracted text is empty.")

    # Metadata
    word_count = len(cleaned_text.split())
    char_count = len(cleaned_text)

    return {
        "text": cleaned_text,
        "word_count": word_count,
        "char_count": char_count
    }
